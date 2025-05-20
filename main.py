import tkinter as tk
from tkinter import messagebox
from tkinter import filedialog
from openpyxl import Workbook
from docx import Document
from recipe_calculator.burger import calculate_burger
from recipe_calculator.pizza import calculate_pizza
from recipe_calculator.wok import calculate_wok

def save_to_doc(calories, price, recipe_name):
    doc = Document()
    doc.add_heading(f'Recipe: {recipe_name}', level=1)
    doc.add_paragraph(f'Calories: {calories}')
    doc.add_paragraph(f'Price: {price}')
    file_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                               filetypes=[("Word Documents", "*.docx")])
    if file_path:
        doc.save(file_path)

def save_to_xls(calories, price, recipe_name):
    wb = Workbook()
    ws = wb.active
    ws.append(['Recipe', 'Calories', 'Price'])
    ws.append([recipe_name, calories, price])
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                               filetypes=[("Excel Files", "*.xlsx")])
    if file_path:
        wb.save(file_path)

def calculate_recipe():
    recipe_name = recipe_var.get()
    ingredients = [
        {'name': 'Ingredient1', 'calories': 100, 'price': 1.5},
        {'name': 'Ingredient2', 'calories': 200, 'price': 2.0},
    ]
    
    if recipe_name == "Burger":
        calories, price = calculate_burger(ingredients)
    elif recipe_name == "Pizza":
        calories, price = calculate_pizza(ingredients)
    elif recipe_name == "Wok":
        calories, price = calculate_wok(ingredients)
    else:
        messagebox.showerror("Error", "Invalid recipe selected.")
        return
    
    result_label.config(text=f'Calories: {calories}, Price: {price:.2f}')
    save_to_doc(calories, price, recipe_name)
    save_to_xls(calories, price, recipe_name)

root = tk.Tk()
root.title("Recipe Calculator")

recipe_var = tk.StringVar(value="Burger")
recipe_label = tk.Label(root, text="Select Recipe:")
recipe_label.pack()

recipe_menu = tk.OptionMenu(root, recipe_var, "Burger", "Pizza", "Wok")
recipe_menu.pack()

calculate_button = tk.Button(root, text="Calculate", command=calculate_recipe)
calculate_button.pack()

result_label = tk.Label(root, text="")
result_label.pack()

root.mainloop()
